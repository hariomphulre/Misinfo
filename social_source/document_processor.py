"""
Document and File Processing System
Handles various file types: PDFs, DOCs, images, videos, etc.
"""

import os
import requests
from pathlib import Path
import mimetypes
import json
from dotenv import load_dotenv
import logging

# For document processing
try:
    import PyPDF2
    import docx
    from PIL import Image
    import cv2
except ImportError:
    print("Some libraries not installed. Install with:")
    print("pip install PyPDF2 python-docx Pillow opencv-python")

load_dotenv()
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

API_BASE_URL = os.getenv("API_BASE_URL", "https://misinformation-collector-zda54hwita-el.a.run.app")

class DocumentProcessor:
    def __init__(self):
        self.supported_types = {
            'application/pdf': self.process_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self.process_docx,
            'application/msword': self.process_doc,
            'text/plain': self.process_text,
            'image/jpeg': self.process_image,
            'image/png': self.process_image,
            'image/gif': self.process_image,
            'video/mp4': self.process_video,
            'video/avi': self.process_video,
            'video/mov': self.process_video
        }

    def process_file(self, file_path, source="file_upload"):
        """Process any supported file type"""
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return None

        file_path = Path(file_path)
        mime_type, _ = mimetypes.guess_type(str(file_path))
        
        if mime_type not in self.supported_types:
            logger.warning(f"Unsupported file type: {mime_type}")
            return self.process_unknown_file(file_path)

        try:
            processor = self.supported_types[mime_type]
            result = processor(file_path)
            
            if result:
                # Send to backend
                self.send_to_backend(result, source)
                return result
            
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {e}")
            return None

    def process_pdf(self, file_path):
        """Extract text from PDF files"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_content = ""
                
                for page_num, page in enumerate(pdf_reader.pages):
                    text_content += page.extract_text() + "\n"

                return {
                    "type": "pdf_document",
                    "content": text_content.strip(),
                    "metadata": {
                        "filename": file_path.name,
                        "file_size": file_path.stat().st_size,
                        "pages": len(pdf_reader.pages),
                        "file_path": str(file_path),
                        "mime_type": "application/pdf"
                    }
                }
        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            return None

    def process_docx(self, file_path):
        """Extract text from DOCX files"""
        try:
            doc = docx.Document(file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"

            return {
                "type": "word_document",
                "content": text_content.strip(),
                "metadata": {
                    "filename": file_path.name,
                    "file_size": file_path.stat().st_size,
                    "paragraphs": len(doc.paragraphs),
                    "file_path": str(file_path),
                    "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                }
            }
        except Exception as e:
            logger.error(f"Error processing DOCX: {e}")
            return None

    def process_doc(self, file_path):
        """Process older DOC files"""
        # This would require python-docx2txt or similar
        logger.warning("DOC files require additional libraries. Converting to text...")
        return self.process_text(file_path)

    def process_text(self, file_path):
        """Process plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()

            return {
                "type": "text_document",
                "content": content,
                "metadata": {
                    "filename": file_path.name,
                    "file_size": file_path.stat().st_size,
                    "encoding": "utf-8",
                    "file_path": str(file_path),
                    "mime_type": "text/plain"
                }
            }
        except Exception as e:
            logger.error(f"Error processing text file: {e}")
            return None

    def process_image(self, file_path):
        """Process image files and extract metadata"""
        try:
            with Image.open(file_path) as img:
                # Extract EXIF data if available
                exif_data = {}
                if hasattr(img, '_getexif'):
                    exif = img._getexif()
                    if exif:
                        exif_data = {str(k): str(v) for k, v in exif.items()}

                return {
                    "type": "image_file",
                    "content": f"Image file: {file_path.name}",
                    "metadata": {
                        "filename": file_path.name,
                        "file_size": file_path.stat().st_size,
                        "dimensions": f"{img.width}x{img.height}",
                        "format": img.format,
                        "mode": img.mode,
                        "exif_data": exif_data,
                        "file_path": str(file_path),
                        "mime_type": f"image/{img.format.lower()}"
                    }
                }
        except Exception as e:
            logger.error(f"Error processing image: {e}")
            return None

    def process_video(self, file_path):
        """Process video files and extract metadata"""
        try:
            cap = cv2.VideoCapture(str(file_path))
            
            # Get video properties
            fps = cap.get(cv2.CAP_PROP_FPS)
            frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
            duration = frame_count / fps if fps > 0 else 0
            width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
            height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
            
            cap.release()

            return {
                "type": "video_file",
                "content": f"Video file: {file_path.name}",
                "metadata": {
                    "filename": file_path.name,
                    "file_size": file_path.stat().st_size,
                    "duration_seconds": duration,
                    "fps": fps,
                    "frame_count": frame_count,
                    "dimensions": f"{width}x{height}",
                    "file_path": str(file_path),
                    "mime_type": mimetypes.guess_type(str(file_path))[0]
                }
            }
        except Exception as e:
            logger.error(f"Error processing video: {e}")
            return None

    def process_unknown_file(self, file_path):
        """Handle unknown file types"""
        try:
            return {
                "type": "unknown_file",
                "content": f"Unknown file type: {file_path.name}",
                "metadata": {
                    "filename": file_path.name,
                    "file_size": file_path.stat().st_size,
                    "file_extension": file_path.suffix,
                    "file_path": str(file_path),
                    "mime_type": mimetypes.guess_type(str(file_path))[0] or "unknown"
                }
            }
        except Exception as e:
            logger.error(f"Error processing unknown file: {e}")
            return None

    def send_to_backend(self, data, source="document_processor"):
        """Send processed document data to backend"""
        try:
            payload = {
                "source": source,
                "type": data.get("type", "document"),
                "content_text": data.get("content", ""),
                "metadata": json.dumps(data.get("metadata", {}))
            }
            
            response = requests.post(f"{API_BASE_URL}/collect", data=payload, timeout=30)
            
            if response.status_code == 200:
                result = response.json()
                logger.info(f"Document sent to backend successfully: {result.get('doc_id')}")
                return result
            else:
                logger.error(f"Failed to send document to backend: {response.status_code}")
                return None
                
        except Exception as e:
            logger.error(f"Error sending document to backend: {e}")
            return None

    def batch_process_directory(self, directory_path, source="batch_upload"):
        """Process all files in a directory"""
        directory = Path(directory_path)
        
        if not directory.exists():
            logger.error(f"Directory not found: {directory_path}")
            return []

        results = []
        for file_path in directory.rglob('*'):
            if file_path.is_file():
                logger.info(f"Processing: {file_path}")
                result = self.process_file(file_path, source)
                if result:
                    results.append(result)

        logger.info(f"Batch processing complete. Processed {len(results)} files.")
        return results

def main():
    """Example usage of Document Processor"""
    processor = DocumentProcessor()
    
    # Example: Process a single file
    # result = processor.process_file("path/to/document.pdf")
    
    # Example: Batch process a directory
    # results = processor.batch_process_directory("path/to/documents/")
    
    print("Document Processor initialized. Use process_file() or batch_process_directory() methods.")

if __name__ == "__main__":
    main()
